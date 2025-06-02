"""
Funções utilitárias: cálculos matemáticos, manipulação de arquivos Word e outras utilidades
"""

import os
import math
import random
import string
import logging
from typing import Any, List, Dict
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from ..exceptions import ToolError
from ...logs import setup_logging

logger = setup_logging(__name__, 'utils.log')

def calcular(expressao: str) -> str:
    """
    Avalia uma expressão matemática de forma segura.
    Args:
        expressao: Expressão matemática como string
    Returns:
        Resultado como string
    """
    try:
        # Limita o escopo para funções matemáticas
        permitido = {k: getattr(math, k) for k in dir(math) if not k.startswith("_")}
        permitido["abs"] = abs
        permitido["round"] = round
        resultado = eval(expressao, {"__builtins__": {}}, permitido)
        return str(resultado)
    except Exception as e:
        logger.error(f"Erro ao calcular expressão: {str(e)}")
        raise ToolError(f"Erro ao calcular expressão: {str(e)}")

# Manipulação básica de arquivos Word (docx)
def criar_word(conteudo: str, caminho: str) -> str:
    """
    Cria um arquivo Word (.docx) com o conteúdo fornecido.
    Args:
        conteudo: Texto a ser inserido
        caminho: Caminho do arquivo a ser salvo
    Returns:
        Mensagem de confirmação
    """
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph(conteudo)
        doc.save(caminho)
        return f"Arquivo Word criado com sucesso: {caminho}"
    except Exception as e:
        logger.error(f"Erro ao criar arquivo Word: {str(e)}")
        raise ToolError(f"Erro ao criar arquivo Word: {str(e)}")

def ler_word(caminho: str) -> str:
    """
    Lê o texto de um arquivo Word (.docx).
    Args:
        caminho: Caminho do arquivo
    Returns:
        Texto extraído
    """
    try:
        from docx import Document
        doc = Document(caminho)
        texto = '\n'.join([p.text for p in doc.paragraphs])
        return texto
    except Exception as e:
        logger.error(f"Erro ao ler arquivo Word: {str(e)}")
        raise ToolError(f"Erro ao ler arquivo Word: {str(e)}")

def gerar_relatorio_word(dados: dict, caminho: str, titulo: str = "Relatório") -> str:
    """
    Gera um relatório Word (.docx) a partir de dados estruturados (dicionário).
    Cria um título, uma tabela com os dados e salva o arquivo.
    Args:
        dados: Dicionário com os dados (chave: lista de valores ou lista de dicionários)
        caminho: Caminho do arquivo a ser salvo
        titulo: Título do relatório
    Returns:
        Mensagem de confirmação
    """
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(titulo, 0)
        
        # Se dados for lista de dicionários
        if isinstance(dados, list) and all(isinstance(item, dict) for item in dados):
            chaves = list(dados[0].keys())
            tabela = doc.add_table(rows=1, cols=len(chaves))
            hdr_cells = tabela.rows[0].cells
            for i, chave in enumerate(chaves):
                hdr_cells[i].text = str(chave)
            for item in dados:
                row_cells = tabela.add_row().cells
                for i, chave in enumerate(chaves):
                    row_cells[i].text = str(item.get(chave, ""))
        # Se dados for dict de listas
        elif isinstance(dados, dict):
            chaves = list(dados.keys())
            num_linhas = len(next(iter(dados.values())))
            tabela = doc.add_table(rows=1, cols=len(chaves))
            hdr_cells = tabela.rows[0].cells
            for i, chave in enumerate(chaves):
                hdr_cells[i].text = str(chave)
            for i in range(num_linhas):
                row_cells = tabela.add_row().cells
                for j, chave in enumerate(chaves):
                    row_cells[j].text = str(dados[chave][i])
        else:
            doc.add_paragraph("Formato de dados não suportado para tabela.")
        
        doc.save(caminho)
        return f"Relatório Word gerado com sucesso: {caminho}"
    except Exception as e:
        logger.error(f"Erro ao gerar relatório Word: {str(e)}")
        raise ToolError(f"Erro ao gerar relatório Word: {str(e)}")

def adicionar_grafico_word(dados: dict, caminho: str, titulo: str = "Gráfico") -> str:
    """
    Gera um gráfico simples (barra) a partir de dados e insere em um arquivo Word.
    Args:
        dados: Dicionário {categoria: valor}
        caminho: Caminho do arquivo Word a ser salvo
        titulo: Título do gráfico
    Returns:
        Mensagem de confirmação
    """
    try:
        import matplotlib.pyplot as plt
        from docx import Document
        from docx.shared import Inches
        import tempfile
        # Gera gráfico
        categorias = list(dados.keys())
        valores = list(dados.values())
        plt.figure(figsize=(6, 4))
        plt.bar(categorias, valores)
        plt.title(titulo)
        plt.tight_layout()
        # Salva imagem temporária
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            plt.savefig(tmp.name)
            plt.close()
            # Adiciona ao Word
            doc = Document()
            doc.add_heading(titulo, 1)
            doc.add_picture(tmp.name, width=Inches(5))
            doc.save(caminho)
        return f"Gráfico adicionado ao Word com sucesso: {caminho}"
    except Exception as e:
        logger.error(f"Erro ao adicionar gráfico ao Word: {str(e)}")
        raise ToolError(f"Erro ao adicionar gráfico ao Word: {str(e)}")

def gerar_senha(tamanho: int = 12) -> str:
    """
    Gera uma senha aleatória segura.
    
    Args:
        tamanho: Tamanho da senha (padrão: 12 caracteres)
        
    Returns:
        Senha gerada
    """
    try:
        # Caracteres permitidos
        letras = string.ascii_letters
        numeros = string.digits
        especiais = "!@#$%^&*()_+-=[]{}|;:,.<>?"
        
        # Garante pelo menos um de cada tipo
        senha = [
            random.choice(letras),
            random.choice(numeros),
            random.choice(especiais)
        ]
        
        # Completa o resto da senha
        todos_caracteres = letras + numeros + especiais
        senha.extend(random.choice(todos_caracteres) for _ in range(tamanho - 3))
        
        # Embaralha a senha
        random.shuffle(senha)
        
        return ''.join(senha)
        
    except Exception as e:
        logger.error(f"Erro ao gerar senha: {str(e)}")
        raise ToolError(f"Erro ao gerar senha: {str(e)}")

def converter_data(data: str, formato_entrada: str, formato_saida: str) -> str:
    """
    Converte uma data de um formato para outro.
    
    Args:
        data: String da data
        formato_entrada: Formato da data de entrada (ex: "%d/%m/%Y")
        formato_saida: Formato da data de saída (ex: "%Y-%m-%d")
        
    Returns:
        Data convertida
    """
    try:
        data_obj = datetime.strptime(data, formato_entrada)
        return data_obj.strftime(formato_saida)
    except Exception as e:
        logger.error(f"Erro ao converter data: {str(e)}")
        raise ToolError(f"Erro ao converter data: {str(e)}")

def calcular_idade(data_nascimento: str, formato: str = "%d/%m/%Y") -> int:
    """
    Calcula a idade a partir da data de nascimento.
    
    Args:
        data_nascimento: Data de nascimento
        formato: Formato da data (padrão: "%d/%m/%Y")
        
    Returns:
        Idade em anos
    """
    try:
        nascimento = datetime.strptime(data_nascimento, formato)
        hoje = datetime.now()
        idade = hoje.year - nascimento.year
        
        # Ajusta idade se ainda não fez aniversário
        if (hoje.month, hoje.day) < (nascimento.month, nascimento.day):
            idade -= 1
            
        return idade
        
    except Exception as e:
        logger.error(f"Erro ao calcular idade: {str(e)}")
        raise ToolError(f"Erro ao calcular idade: {str(e)}")

def enviar_email(destinatario: str, assunto: str, corpo: str) -> str:
    """
    Envia um e-mail usando SMTP.
    
    Args:
        destinatario: E-mail do destinatário
        assunto: Assunto do e-mail
        corpo: Corpo do e-mail
        
    Returns:
        Mensagem de confirmação
    """
    try:
        # Configurações do servidor SMTP (exemplo)
        smtp_server = os.getenv("SMTP_SERVER", "smtp.gmail.com")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))
        smtp_user = os.getenv("SMTP_USER", "")
        smtp_password = os.getenv("SMTP_PASSWORD", "")
        
        # Cria mensagem
        msg = MIMEMultipart()
        msg['From'] = smtp_user
        msg['To'] = destinatario
        msg['Subject'] = assunto
        msg.attach(MIMEText(corpo, 'plain'))
        
        # Envia e-mail
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
            
        return f"E-mail enviado com sucesso para {destinatario}"
        
    except Exception as e:
        logger.error(f"Erro ao enviar e-mail: {str(e)}")
        raise ToolError(f"Erro ao enviar e-mail: {str(e)}")

def gerar_relatorio(dados: List[Dict[str, Any]]) -> str:
    """
    Gera um relatório formatado a partir de uma lista de dicionários.
    
    Args:
        dados: Lista de dicionários com os dados
        
    Returns:
        Relatório formatado
    """
    try:
        if not dados:
            return "Nenhum dado para gerar relatório"
            
        # Obtém cabeçalhos
        cabecalhos = list(dados[0].keys())
        
        # Calcula largura das colunas
        larguras = {h: len(str(h)) for h in cabecalhos}
        for linha in dados:
            for h in cabecalhos:
                larguras[h] = max(larguras[h], len(str(linha.get(h, ""))))
        
        # Gera relatório
        relatorio = []
        
        # Cabeçalho
        linha_cabecalho = " | ".join(f"{h:{larguras[h]}}" for h in cabecalhos)
        relatorio.append(linha_cabecalho)
        relatorio.append("-" * len(linha_cabecalho))
        
        # Dados
        for linha in dados:
            linha_dados = " | ".join(f"{str(linha.get(h, '')):{larguras[h]}}" for h in cabecalhos)
            relatorio.append(linha_dados)
            
        return "\n".join(relatorio)
        
    except Exception as e:
        logger.error(f"Erro ao gerar relatório: {str(e)}")
        raise ToolError(f"Erro ao gerar relatório: {str(e)}") 