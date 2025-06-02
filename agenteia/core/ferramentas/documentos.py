"""
Ferramentas para manipulação de documentos
"""

import os
from typing import List, Dict, Any, Optional
from datetime import datetime
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import json

def criar_documento_word(
    titulo: str,
    conteudo: List[Dict[str, Any]],
    caminho_saida: str,
    estilo: str = "padrao"
) -> str:
    """
    Cria um documento Word com formatação profissional.
    
    Args:
        titulo: Título do documento
        conteudo: Lista de seções com título e texto
        caminho_saida: Caminho para salvar o documento
        estilo: Estilo do documento (padrao, relatorio, curriculo)
        
    Returns:
        Caminho do arquivo criado
    """
    try:
        # Criar documento
        doc = docx.Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Adicionar título
        titulo_paragrafo = doc.add_paragraph()
        titulo_run = titulo_paragrafo.add_run(titulo)
        titulo_run.bold = True
        titulo_run.font.size = Pt(16)
        titulo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar data
        data_paragrafo = doc.add_paragraph()
        data_run = data_paragrafo.add_run(datetime.now().strftime("%d/%m/%Y"))
        data_paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Adicionar conteúdo
        for secao in conteudo:
            # Título da seção
            if "titulo" in secao:
                titulo_secao = doc.add_paragraph()
                titulo_secao_run = titulo_secao.add_run(secao["titulo"])
                titulo_secao_run.bold = True
                titulo_secao_run.font.size = Pt(14)
            
            # Texto da seção
            if "texto" in secao:
                texto_paragrafo = doc.add_paragraph()
                texto_paragrafo.add_run(secao["texto"])
            
            # Lista de itens
            if "itens" in secao:
                for item in secao["itens"]:
                    item_paragrafo = doc.add_paragraph(style='List Bullet')
                    item_paragrafo.add_run(str(item))
            
            # Tabela
            if "tabela" in secao:
                tabela = doc.add_table(rows=1, cols=len(secao["tabela"][0]))
                tabela.style = 'Table Grid'
                
                # Cabeçalho
                header_cells = tabela.rows[0].cells
                for i, header in enumerate(secao["tabela"][0]):
                    header_cells[i].text = str(header)
                
                # Dados
                for row_data in secao["tabela"][1:]:
                    row_cells = tabela.add_row().cells
                    for i, cell_data in enumerate(row_data):
                        row_cells[i].text = str(cell_data)
        
        # Salvar documento
        doc.save(caminho_saida)
        return caminho_saida
        
    except Exception as e:
        raise Exception(f"Erro ao criar documento Word: {e}")

def criar_curriculo(
    dados: Dict[str, Any],
    caminho_saida: str
) -> str:
    """
    Cria um currículo profissional em formato Word.
    
    Args:
        dados: Dicionário com dados do currículo
        caminho_saida: Caminho para salvar o documento
        
    Returns:
        Caminho do arquivo criado
    """
    try:
        conteudo = [
            {
                "titulo": "Dados Pessoais",
                "texto": f"{dados.get('nome', '')}\n{dados.get('email', '')}\n{dados.get('telefone', '')}\n{dados.get('endereco', '')}"
            },
            {
                "titulo": "Objetivo",
                "texto": dados.get('objetivo', '')
            },
            {
                "titulo": "Formação Acadêmica",
                "itens": dados.get('formacao', [])
            },
            {
                "titulo": "Experiência Profissional",
                "itens": dados.get('experiencia', [])
            },
            {
                "titulo": "Habilidades",
                "itens": dados.get('habilidades', [])
            }
        ]
        
        return criar_documento_word("Currículo", conteudo, caminho_saida, "curriculo")
        
    except Exception as e:
        raise Exception(f"Erro ao criar currículo: {e}")

def criar_relatorio(
    titulo: str,
    dados: Dict[str, Any],
    caminho_saida: str
) -> str:
    """
    Cria um relatório profissional em formato Word.
    
    Args:
        titulo: Título do relatório
        dados: Dicionário com dados do relatório
        caminho_saida: Caminho para salvar o documento
        
    Returns:
        Caminho do arquivo criado
    """
    try:
        conteudo = [
            {
                "titulo": "Resumo Executivo",
                "texto": dados.get('resumo', '')
            },
            {
                "titulo": "Metodologia",
                "texto": dados.get('metodologia', '')
            },
            {
                "titulo": "Resultados",
                "texto": dados.get('resultados', '')
            },
            {
                "titulo": "Análise de Dados",
                "tabela": dados.get('tabela_dados', [])
            },
            {
                "titulo": "Conclusões",
                "itens": dados.get('conclusoes', [])
            },
            {
                "titulo": "Recomendações",
                "itens": dados.get('recomendacoes', [])
            }
        ]
        
        return criar_documento_word(titulo, conteudo, caminho_saida, "relatorio")
        
    except Exception as e:
        raise Exception(f"Erro ao criar relatório: {e}")

def converter_para_word(
    arquivo_entrada: str,
    caminho_saida: str
) -> str:
    """
    Converte arquivos de outros formatos para Word.
    
    Args:
        arquivo_entrada: Caminho do arquivo de entrada
        caminho_saida: Caminho para salvar o documento Word
        
    Returns:
        Caminho do arquivo Word criado
    """
    try:
        # Detectar extensão
        _, extensao = os.path.splitext(arquivo_entrada)
        
        if extensao.lower() == '.csv':
            # Converter CSV para Word
            df = pd.read_csv(arquivo_entrada)
            conteudo = [
                {
                    "titulo": "Dados do CSV",
                    "tabela": [df.columns.tolist()] + df.values.tolist()
                }
            ]
            return criar_documento_word("Relatório CSV", conteudo, caminho_saida)
            
        elif extensao.lower() == '.json':
            # Converter JSON para Word
            with open(arquivo_entrada, 'r', encoding='utf-8') as f:
                dados = json.load(f)
            conteudo = [
                {
                    "titulo": "Dados do JSON",
                    "texto": json.dumps(dados, indent=2, ensure_ascii=False)
                }
            ]
            return criar_documento_word("Relatório JSON", conteudo, caminho_saida)
            
        else:
            raise Exception(f"Formato não suportado: {extensao}")
            
    except Exception as e:
        raise Exception(f"Erro ao converter para Word: {e}") 